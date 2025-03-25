from lxml import etree # type: ignore
from docx import Document # type: ignore

with open('Src/arquivo.xml', 'rb') as file:
    xml_data = file.read()

root = etree.fromstring(xml_data)

namespaces = {'nfe': 'http://www.portalfiscal.inf.br/nfe'}
# Definir namespaces e buscar dados do XML
namespaces = {'nfe': 'http://www.portalfiscal.inf.br/nfe'}
item_element = root.xpath('//nfe:ide/nfe:cNF', namespaces=namespaces)
nNF_element = root.xpath('//nfe:ide/nfe:nNF', namespaces=namespaces)
cProd_element = root.xpath('//nfe:prod/nfe:cProd', namespaces=namespaces)
xProd_element = root.xpath('//nfe:prod/nfe:xProd', namespaces=namespaces)
ncm_elements = root.xpath('//nfe:det/nfe:prod/nfe:NCM', namespaces=namespaces)
qtrib_elements = root.xpath('//nfe:det/nfe:prod/nfe:qTrib', namespaces=namespaces)
cfop_elements = root.xpath('//nfe:det/nfe:prod/nfe:CFOP', namespaces=namespaces)


# Extrair valores do XML
item_value = item_element[0].text if item_element else "Valor não encontrado"
nNF_value = nNF_element[0].text if nNF_element else "Valor não encontrado"
cProd_value = cProd_element[0].text if cProd_element else "Valor não encontrado"
xProd_value = xProd_element[0].text if xProd_element else "Valor não encontrado"
ncm_value = ncm_elements[0].text if ncm_elements else "Valor não encontrado"
qtrib_value = qtrib_elements[0].text if qtrib_elements else "Valor não encontrado"
cfop_value = cfop_elements[0].text if cfop_elements else "Valor não encontrado"


doc = Document(r'C:\Users\EricCosta\Documents\Projeto_Mercato\NFC-maneger\Src\CertificadoConf.docx')

# Substituir ou adicionar o texto acima da tabela
for paragraph in doc.paragraphs:
    if "NF" in paragraph.text:  # Localiza o parágrafo com "NF: N°"
        formatted_value = "{:09}".format(int(nNF_value))  # Formata para 000.000.000
        formatted_value = f"{formatted_value[:3]}.{formatted_value[3:6]}.{formatted_value[6:]}"  # Adiciona os pontos
        paragraph.text = f"NF: N° {formatted_value}"  # Substitui pelo valor de <nNF>

# Parte para os dados nas tabelas
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "Código" in cell.text:  # Localizar a célula que contém "Código"
                # Preenche apenas na próxima linha da tabela
                for item_row in table.rows[1:]:  # Evitar a linha de cabeçalho
                    item_row.cells[0].text = cProd_value
                    break  # Inserir apenas uma vez e sair

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
                # Preenche apenas na próxima linha da tabela
                for item_row in table.rows[1:]:  # Evitar a linha de cabeçalho
                    item_row.cells[1].text = xProd_value
                    break  # Inserir apenas uma vez e sair

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
                # Preenche apenas na próxima linha da tabela
                for item_row in table.rows[1:]:  # Evitar a linha de cabeçalho
                    item_row.cells[2].text = qtrib_value
                    break  # Inserir apenas uma vez e sair

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
                for item_row in table.rows[1:]:  # Evitar a linha de cabeçalho
                    item_row.cells[3].text = ncm_value
                    break  # Inserir apenas uma vez e sair

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
                for item_row in table.rows[1:]:  # Evitar a linha de cabeçalho
                    item_row.cells[4].text = cfop_value
                    break  # Inserir apenas uma vez e sair

doc.save('arquivo_atualizado.docx')